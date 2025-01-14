from datetime import datetime
import openpyxl
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

def format_date(date):
    """
    Đảm bảo định dạng ngày là DD/MM/YYYY.
    :param date: Đối tượng datetime hoặc chuỗi ngày.
    :return: Chuỗi ngày ở định dạng DD/MM/YYYY.
    """
    if isinstance(date, datetime):  # Nếu là đối tượng datetime
        return date.strftime("%d/%m/%Y")
    elif isinstance(date, str):  # Nếu là chuỗi
        try:
            # Xử lý chuỗi có thể ở dạng YYYY-MM-DD
            return datetime.strptime(date, "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            pass
    return date  # Trả về giá trị ban đầu nếu không hợp lệ


def fill_cccd_to_table(doc, cccd):
    """
    Điền từng ký tự của CCCD vào bảng Word với căn giữa và spacing 0.07 inch.
    :param doc: Document object của Word.
    :param cccd: Chuỗi CCCD (12 ký tự).
    """
    cccd = str(cccd)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) - 1 == len(cccd):
                for idx, char in enumerate(cccd):
                    cell = row.cells[idx + 1]

                    # Xóa nội dung hiện tại
                    cell._element.clear_content()

                    # Tạo paragraph mới
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Thiết lập spacing trước và sau (0.07 inch = 101 twips)
                    pPr = paragraph._p.get_or_add_pPr()
                    spacing = parse_xml(r'<w:spacing {} w:before="101" w:after="101"/>'.format(nsdecls('w')))
                    pPr.append(spacing)

                    # Thêm ký tự với định dạng
                    run = paragraph.add_run(char)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

                    # Căn giữa theo chiều dọc
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
                    tcPr.append(tcVAlign)
                break


def fill_specific_cell(table, keyword, value):
    """
    Điền giá trị vào một ô cụ thể trong bảng dựa trên từ khóa,
    giữ nguyên định dạng bảng và kích thước các ô.
    :param table: Bảng trong Word.
    :param keyword: Từ khóa để xác định ô (nằm trong cùng hàng với ô cần điền).
    :param value: Giá trị cần điền (thay thế cho từ khóa).
    """
    for row in table.rows:
        for cell in row.cells:
            if keyword in cell.text:
                # Lưu lại thông tin về kích thước ô
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Lưu chiều rộng của ô
                grid_span = tcPr.find(qn('w:gridSpan'))
                width = tcPr.find(qn('w:tcW'))

                # Thay thế nội dung trong paragraph chứa từ khóa
                for paragraph in cell.paragraphs:
                    if keyword in paragraph.text:
                        # Lưu lại định dạng paragraph
                        original_alignment = paragraph.alignment
                        original_spacing_before = None
                        original_spacing_after = None

                        # Lưu spacing nếu có
                        if paragraph._p.pPr is not None:
                            spacing_element = paragraph._p.pPr.spacing
                            if spacing_element is not None:
                                original_spacing_before = spacing_element.get(qn('w:before'))
                                original_spacing_after = spacing_element.get(qn('w:after'))

                        # Thay thế text
                        new_text = paragraph.text.replace(keyword, str(value) if value else "")

                        # Xóa runs hiện tại
                        p = paragraph._p
                        for run in paragraph.runs:
                            p.remove(run._r)

                        # Thêm run mới
                        run = paragraph.add_run(new_text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

                        # Khôi phục định dạng paragraph
                        paragraph.alignment = original_alignment

                        # Khôi phục spacing nếu có
                        if original_spacing_before or original_spacing_after:
                            pPr = paragraph._p.get_or_add_pPr()
                            spacing = parse_xml(
                                r'<w:spacing {} w:before="{}" w:after="{}"/>'.format(
                                    nsdecls('w'),
                                    original_spacing_before or "0",
                                    original_spacing_after or "0"
                                )
                            )
                            pPr.append(spacing)

                # Khôi phục chiều rộng và gridSpan của ô
                if width is not None:
                    new_width = parse_xml(width.xml)
                    tcPr.append(new_width)
                if grid_span is not None:
                    new_grid_span = parse_xml(grid_span.xml)
                    tcPr.append(new_grid_span)

                return

def fill_word_template_with_table(excel_file, word_template, output_folder):
    workbook = openpyxl.load_workbook(excel_file)
    sheet = workbook.active
    headers = [cell.value for cell in sheet[1]]

    for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        doc = Document(word_template)
        data = dict(zip(headers, row))

        ho_ten_value = data.get("HỌ TÊN", "")
        if ho_ten_value:
            data["HỌ TÊN"] = ho_ten_value.upper()  # Phiên bản in hoa để dùng ngoài table
            data["HỌ TÊN2"] = ho_ten_value  # Phiên bản giữ nguyên để dùng trong table

        #Skip empty rows or rows with "HỌ TÊN = Phòng trống"
        if not ho_ten_value or ho_ten_value == "Phòng trống":
            continue

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key == "NGÀY SINH" and value:
                    value = format_date(value)
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value) if value else "")

        if "SỐ CCCD" in data and data["SỐ CCCD"]:
            fill_cccd_to_table(doc, data["SỐ CCCD"])

        if "HỌ TÊN2" in data:
                for table in doc.tables:
                    fill_specific_cell(table, "HỌ TÊN", data["HỌ TÊN2"])

        output_path = f"{output_folder}/output_{row_idx - 1}.docx"
        doc.save(output_path)
        print(f"Document saved: {output_path}")



excel_file = "input.xlsx" 
word_template = "template.docx"
output_folder = "output_docs" 

